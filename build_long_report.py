import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os

doc = docx.Document()

# 1. Page Setup
for section in doc.sections:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.98)     # ~2.5 cm
    section.bottom_margin = Inches(0.98)  # ~2.5 cm
    section.left_margin = Inches(1.18)    # ~3.0 cm
    section.right_margin = Inches(0.79)   # ~2.0 cm

style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(13)
style_normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
style_normal.paragraph_format.line_spacing = 1.3
style_normal.paragraph_format.space_after = Pt(6)
style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    return p

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13.5)
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    return p

def add_heading_1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    return p

def add_heading_2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13.5)
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0x89)
    return p

def add_heading_3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    return p

def add_p(text, bold_prefix="", italic=False):
    p = doc.add_paragraph()
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
    run_t = p.add_run(text)
    run_t.italic = italic
    return p

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
    p.add_run(text)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    shd = parse_xml(r'<w:shd {} w:fill="F3F4F6"/>'.format(nsdecls('w')))
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    return p

def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for cell in hdr_cells:
        shd = parse_xml(r'<w:shd {} w:fill="1E3A8A"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shd)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
    for r_idx in range(1, len(table.rows)):
        row = table.rows[r_idx]
        bg_color = "F9FAFB" if r_idx % 2 == 1 else "FFFFFF"
        for cell in row.cells:
            shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
            cell._tc.get_or_add_tcPr().append(shd)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(10)

def main():
    import chapter1
    import chapter2
    import chapter3
    import chapter4
    import chapter5
    import chapter6
    import chapter7
    import chapter8_refs

    # BÌA
    add_title("BÁO CÁO BÀI TẬP LỚN HỌC PHẦN\nXỬ LÝ NGÔN NGỮ TỰ NHIÊN")
    add_subtitle("Đề tài: Hệ thống Tóm tắt Văn bản Trích xuất Song ngữ (English & Vietnamese) dựa trên Quy trình 2 Giai đoạn: SBERT Fine-Tuned, K-Means Selection và Post-Filtering")
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(28)
    r_m = p_meta.add_run("Mô hình Đề xuất: FineTuned-SBERT-KMeans | Dữ liệu Thực nghiệm: VietNews & CNN/DailyMail\nPublic Model Hub: kttt294/vietnamese-sbert-finetuned | Triển khai: FastAPI, React Vite, Docker Compose")
    r_m.italic = True
    r_m.font.size = Pt(11)
    r_m.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    doc.add_page_break()

    chapter1.write(doc, add_heading_1, add_heading_2, add_heading_3, add_p, add_bullet, add_code, style_table)
    doc.add_page_break()
    chapter2.write(doc, add_heading_1, add_heading_2, add_heading_3, add_p, add_bullet, add_code, style_table)
    doc.add_page_break()
    chapter3.write(doc, add_heading_1, add_heading_2, add_heading_3, add_p, add_bullet, add_code, style_table)
    doc.add_page_break()
    chapter4.write(doc, add_heading_1, add_heading_2, add_heading_3, add_p, add_bullet, add_code, style_table)
    doc.add_page_break()
    chapter5.write(doc, add_heading_1, add_heading_2, add_heading_3, add_p, add_bullet, add_code, style_table)
    doc.add_page_break()
    chapter6.write(doc, add_heading_1, add_heading_2, add_heading_3, add_p, add_bullet, add_code, style_table)
    doc.add_page_break()
    chapter7.write(doc, add_heading_1, add_heading_2, add_heading_3, add_p, add_bullet, add_code, style_table)
    doc.add_page_break()
    chapter8_refs.write(doc, add_heading_1, add_heading_2, add_heading_3, add_p, add_bullet, add_code, style_table)

    output_path = r"C:\Users\trang\Desktop\NLP_BTL\BaoCao_HoanChinh_Long.docx"
    doc.save(output_path)
    print("Saved to", output_path)

if __name__ == '__main__':
    main()
